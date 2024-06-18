import base64
from io import BytesIO
import logging
from typing import Any
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from helpers.file_system_helper import get_required_env

#################################
# from dotenv import load_dotenv
# load_dotenv()
#################################

# Constant values
FONT_FAMILY = "EYInterstate"
EY_YELLOW = RGBColor(0xff, 0xe6, 0x00)
BLACK_COLOR = RGBColor(0x00, 0x00, 0x00)

BULLET_AFTER_SPACE = int(get_required_env("BULLET_AFTER_SPACE"))
PARAGRAPH_AFTER_SPACE = int(get_required_env("PARAGRAPH_AFTER_SPACE"))
HEADING_AFTER_SPACE = int(get_required_env("HEADING_AFTER_SPACE"))
BULLET_FONT_SIZE = int(get_required_env("BULLET_FONT_SIZE"))
PARAGRAPH_FONT_SIZE = int(get_required_env("PARAGRAPH_FONT_SIZE"))
HEADING_FONT_SIZE = int(get_required_env("HEADING_FONT_SIZE"))

def format_text(
        component: Any,
        content: list,
        font_size: int,
):
    try:
        for data in content:
            text = data.get("text", None)
            marks = data.get("marks", [])

            if text:
                component_run = component.add_run(text)
                component_run.bold = False
                component_run.italic = False
                component_run.underline = False
                component_run.font.name = FONT_FAMILY
                component_run.font.color.rgb = BLACK_COLOR
                component_run.font.size = Pt(font_size)

            for mark in marks:
                mark_type = mark.get("type", None)
                if mark_type == "bold":
                    component_run.bold = True
                elif mark_type == "italic":
                    component_run.italic = True
                elif mark_type == "underline":
                    component_run.underline = True
    except Exception as e:
        logging.error(str(e))

def add_heading_to_doc(
        document: Any,
        content: list,
        attrs: dict
):
    try:
        heading_level = int(attrs.get("level", 1))
        heading = document.add_heading(level=heading_level)
        heading.paragraph_format.space_after = Pt(HEADING_AFTER_SPACE)

        modified_content = [{**i, "marks": [{"type": "bold"}]} for i in content]

        format_text(
            component=heading,
            content=modified_content,
            font_size=(HEADING_FONT_SIZE - heading_level)
        )
    
    except Exception as e:
        logging.error(str(e))

def add_paragraph_to_doc(
        document: Any,
        content: list
):
    try:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(PARAGRAPH_AFTER_SPACE)

        format_text(
            component=paragraph,
            content=content,
            font_size=PARAGRAPH_FONT_SIZE
        )
    
    except Exception as e:
        logging.error(str(e))

def add_bullets_to_doc(
        document: Any,
        content: list,
        list_type: str,
        depth: int = 0
):
    try:
        bullet_style = "List Bullet"
        if list_type == "orderedList":
            bullet_style = "List Number"
        if depth:
            bullet_style += f" {min(depth+1, 3)}"

        for list_item in content:
            item_type = list_item.get("type", None)
            item_content = list_item.get("content", [])
            if item_type == "listItem":
                bullet = document.add_paragraph(style=bullet_style)
                for item in item_content:
                    list_item_type = item.get("type", None)
                    list_item_content = item.get("content", [])
                    if list_item_type == "paragraph":
                        bullet.paragraph_format.space_after = Pt(BULLET_AFTER_SPACE)

                        format_text(
                            component=bullet,
                            content=list_item_content,
                            font_size=BULLET_FONT_SIZE
                        )
                    if list_item_type in ["bulletList", "orderedList"]:
                        add_bullets_to_doc(
                            document=document,
                            content=list_item_content,
                            list_type=list_type,
                            depth=depth+1
                        )
    except Exception as e:
        logging.error(str(e))

def add_page_break(
        document: Any
):
    page_break = OxmlElement('w:br')
    page_break.set(qn('w:type'), 'page')

    document.add_paragraph()._element.append(page_break)

def create_document(
        file_content: list
):
    try:
        document = Document()

        # for data_details in file_content:
        #     content_type = data_details.get("type", None)
        #     content_attrs = data_details.get("attrs", {})
        #     content = data_details.get("content", [])
        #     if content_type == "heading":
        #         add_heading_to_doc(
        #             document=document,
        #             content=content,
        #             attrs=content_attrs
        #         )
        #     if content_type == "paragraph":
        #         add_paragraph_to_doc(
        #             document=document,
        #             content=content
        #         )
        #     if content_type in ["bulletList", "orderedList"]:
        #         add_bullets_to_doc(
        #             document=document,
        #             content=content,
        #             list_type=content_type
        #         )

        for index, observation in enumerate(file_content):
            if index != 0:
                add_page_break(
                    document=document
                )
            observation_content = observation.get("content", [])
            for data_in_observation in observation_content:
                content_type = data_in_observation.get("type", None)
                content_attrs = data_in_observation.get("attrs", {})
                content = data_in_observation.get("content", [])
                if content_type == "heading":
                    add_heading_to_doc(
                        document=document,
                        content=content,
                        attrs=content_attrs
                    )
                if content_type == "paragraph":
                    add_paragraph_to_doc(
                        document=document,
                        content=content
                    )
                if content_type in ["bulletList", "orderedList"]:
                    add_bullets_to_doc(
                        document=document,
                        content=content,
                        list_type=content_type
                    )

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return buffer
    except Exception as e:
        logging.error(str(e))
        raise Exception(f"Error creating document: {e}") from e

def doc_to_base64(
        buffer: BytesIO
):
    try:
        base64file = base64.b64encode(buffer.read()).decode('utf-8')
        return base64file
    except Exception as e:
        logging.error(str(e))
        raise Exception(f"Error converting document to base64: {e}") from e
